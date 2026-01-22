"""
Word Document Generator for Email Thread Summaries
"""

import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Gen II footer text to exclude from LLM processing
GEN_II_FOOTER_START = "NOTICE: Unless otherwise stated, the content of this e-mail"
GEN_II_FOOTER_END = "which can be found here."


def strip_gen_ii_footer(content):
    """
    Removes the Gen II legal footer from email content to save tokens.

    Args:
        content: Email content string

    Returns:
        Content with footer removed if found, original content otherwise
    """
    if not content:
        return content

    # Look for the footer start marker
    # Use case-insensitive search and handle variations in whitespace
    footer_pattern = re.compile(
        r"NOTICE:\s*Unless otherwise stated.*?which can be found here\.", re.IGNORECASE | re.DOTALL
    )

    # Try to find and remove the footer
    cleaned = footer_pattern.sub("", content)

    # Also try a more flexible approach - look for just the start phrase
    # and remove everything from there to the end if it's near the end of the content
    start_marker = re.compile(r"NOTICE:\s*Unless otherwise stated", re.IGNORECASE)
    match = start_marker.search(cleaned)

    if match:
        # If the marker appears in the last 30% of the content, it's likely a footer
        marker_pos = match.start()
        content_length = len(cleaned)
        if marker_pos > content_length * 0.7:
            # Remove everything from the marker to the end
            cleaned = cleaned[:marker_pos].rstrip()

    return cleaned


def format_thread_content(thread):
    """
    Formats a thread (list of messages) into a readable string for LLM processing.

    Args:
        thread: List of message dictionaries

    Returns:
        Formatted string containing the thread content
    """
    formatted_lines = []

    # Sort messages by timestamp to ensure chronological order
    # Normalize timestamps to datetime objects for safe sorting
    def normalize_timestamp(msg):
        ts = msg.get("timestamp")
        if ts is None:
            return datetime.min
        if isinstance(ts, datetime):
            return ts
        if isinstance(ts, str):
            # Try to parse string timestamp if needed
            try:
                return datetime.fromisoformat(ts.replace("Z", "+00:00"))
            except (ValueError, AttributeError):
                return datetime.min
        # For other types, fall back to min
        return datetime.min

    sorted_thread = sorted(thread, key=normalize_timestamp)

    for i, msg in enumerate(sorted_thread, 1):
        formatted_lines.append(f"\n--- Message {i} ---")
        formatted_lines.append(f"From: {msg.get('from', 'Unknown')}")
        formatted_lines.append(f"Date: {msg.get('date', 'Unknown')}")
        formatted_lines.append(f"Subject: {msg.get('subject', 'No Subject')}")
        if msg.get("flag_status"):
            formatted_lines.append(f"Flag Status: {msg.get('flag_status')}")
        formatted_lines.append("\nContent:")
        # Strip Gen II footer before adding content to save tokens
        content = strip_gen_ii_footer(msg.get("content", ""))
        formatted_lines.append(content)
        formatted_lines.append("\n" + "=" * 80)

    return "\n".join(formatted_lines)


def create_summary_document(threads_with_summaries, output_path):
    """
    Creates a Word document with email thread summaries.

    Args:
        threads_with_summaries: List of dicts with keys:
            - 'subject': Thread subject
            - 'summary': Generated summary text
            - 'thread': Original thread data (for reference)
        output_path: Path where the Word document should be saved

    Returns:
        Path to the created document, or None if creation failed
    """
    try:
        doc = Document()

        # Title
        title = doc.add_heading("Email Thread Summaries", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para_format = date_para.runs[0].font
        date_para_format.size = Pt(10)
        date_para_format.italic = True

        doc.add_paragraph()  # Spacing

        # Add summary for each thread
        for idx, item in enumerate(threads_with_summaries, 1):
            # Subject line as heading (level 1)
            subject = item.get("subject", "No Subject")
            doc.add_heading(subject, level=1)

            # Client name on new line (regular paragraph, slightly indented)
            client_name = item.get("client_name", "Unknown Client")
            client_para = doc.add_paragraph(client_name)
            client_format = client_para.runs[0].font
            client_format.size = Pt(11)

            # Summary content
            doc.add_paragraph()  # Spacing before summary
            summary_para = doc.add_paragraph(item["summary"])
            summary_format = summary_para.runs[0].font
            summary_format.size = Pt(11)

            # SF Note section
            doc.add_paragraph()  # Spacing before SF Note
            doc.add_heading("SF Note:", level=2)
            sf_note_content = item.get("sf_note", "SF Note not generated.")
            sf_note_para = doc.add_paragraph(sf_note_content)
            sf_note_format = sf_note_para.runs[0].font
            sf_note_format.size = Pt(11)

            # Add separator
            if idx < len(threads_with_summaries):
                doc.add_paragraph()  # Spacing
                doc.add_paragraph("─" * 80)
                doc.add_paragraph()  # Spacing

        # Save document
        doc.save(output_path)
        return output_path

    except Exception as e:
        print(f"Error creating Word document: {e}")
        return None
