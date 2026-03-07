"""Word document generation for email thread summaries."""

from __future__ import annotations

from datetime import datetime
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

if TYPE_CHECKING:
    from outlook_bot.core.models import ThreadSummary


def create_summary_document(summaries: list[ThreadSummary], output_path: str) -> str | None:
    """Create a Word document with email thread summaries.

    Returns the output path on success, or None on failure.
    """
    try:
        doc = Document()

        title = doc.add_heading("Email Thread Summaries", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_fmt = date_para.runs[0].font
        date_fmt.size = Pt(10)
        date_fmt.italic = True

        doc.add_paragraph()

        for idx, item in enumerate(summaries, 1):
            doc.add_heading(item.subject, level=1)

            client_para = doc.add_paragraph(item.client_name)
            client_para.runs[0].font.size = Pt(11)

            doc.add_paragraph()
            summary_para = doc.add_paragraph(item.summary)
            summary_para.runs[0].font.size = Pt(11)

            doc.add_paragraph()
            doc.add_heading("SF Note:", level=2)
            sf_note_para = doc.add_paragraph(item.sf_note or "SF Note not generated.")
            sf_note_para.runs[0].font.size = Pt(11)

            if idx < len(summaries):
                doc.add_paragraph()
                doc.add_paragraph("\u2500" * 80)
                doc.add_paragraph()

        doc.save(output_path)
        return output_path

    except Exception as e:
        print(f"Error creating Word document: {e}")
        return None
